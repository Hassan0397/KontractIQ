"""
KontractIQ - Create Contract Page
Premium contract template creation with real-time validation
"""

import streamlit as st
from datetime import datetime
from typing import Dict, Any, List, Optional
import re

from ..utils.constants import COLORS, TEMPLATES, TEMPLATE_DESCRIPTIONS
from ..core.templates import (
    ContractTemplate,
    NDAUnilateralTemplate,
    NDAMutualTemplate,
    MSATemplate,
    IndependentContractorTemplate,
    EmploymentOfferTemplate,
    SoftwareLicenseTemplate,
    get_all_templates,
    get_template_fields,
    validate_template_data,
    generate_contract
)

# ============================================================================
# SHADOWS - Defined here to avoid import issues
# ============================================================================

SHADOWS = {
    "sm": "0 1px 3px rgba(10, 38, 71, 0.06)",
    "md": "0 4px 12px rgba(10, 38, 71, 0.08)",
    "lg": "0 8px 24px rgba(10, 38, 71, 0.10)",
    "xl": "0 16px 48px rgba(10, 38, 71, 0.12)",
}

# ============================================================================
# MAIN RENDER FUNCTION
# ============================================================================

def render_create():
    """Render the contract creation page with premium UI"""
    
    # Hero Section with Premium Design
    render_hero_section()
    
    # Quick Stats
    render_quick_stats()
    
    # Main Content
    st.divider()
    
    col1, col2 = st.columns([2.2, 0.9], gap="large")
    
    with col1:
        render_template_selector()
    
    with col2:
        render_premium_features()
    
    st.divider()
    
    # Render selected template form
    selected_template = st.session_state.get('selected_template', TEMPLATES[0])
    
    # Template renderer mapping
    template_renderers = {
        "NDA (Unilateral)": render_nda_unilateral,
        "NDA (Mutual)": render_nda_mutual,
        "Master Services Agreement": render_msa,
        "Independent Contractor Agreement": render_ica,
        "Employment Offer Letter": render_employment,
        "Software License Agreement": render_software_license
    }
    
    renderer = template_renderers.get(selected_template)
    if renderer:
        renderer()
    else:
        st.warning("⚠️ Template not found. Please select another template.")
    
    # Footer with tips
    render_footer_tips()


# ============================================================================
# HERO SECTION
# ============================================================================

def render_hero_section():
    """Render premium hero section"""
    
    st.markdown(
        f"""
        <div style="
            background: linear-gradient(135deg, {COLORS['primary']['gradient_start']} 0%, {COLORS['primary']['gradient_end']} 100%);
            padding: 32px 40px;
            border-radius: 24px;
            margin-bottom: 24px;
            box-shadow: {SHADOWS['xl']};
            position: relative;
            overflow: hidden;
        ">
            <div style="position: absolute; top: -50%; right: -10%; font-size: 200px; opacity: 0.05;">
                ⚖️
            </div>
            <div style="display: flex; justify-content: space-between; align-items: center; position: relative; z-index: 1;">
                <div>
                    <h1 style="color: white; font-size: 28px; font-weight: 700; margin: 0; letter-spacing: -0.5px;">
                        📝 Create Contract
                    </h1>
                    <p style="color: rgba(255,255,255,0.8); font-size: 16px; margin: 8px 0 0 0;">
                        Generate professional contracts from intelligent templates with real-time validation
                    </p>
                </div>
                <div style="display: flex; gap: 12px;">
                    <span style="
                        background: rgba(255,255,255,0.15);
                        color: white;
                        padding: 6px 16px;
                        border-radius: 20px;
                        font-size: 12px;
                        font-weight: 500;
                        backdrop-filter: blur(10px);
                    ">
                        🚀 6 Templates Available
                    </span>
                    <span style="
                        background: rgba(255,255,255,0.15);
                        color: white;
                        padding: 6px 16px;
                        border-radius: 20px;
                        font-size: 12px;
                        font-weight: 500;
                        backdrop-filter: blur(10px);
                    ">
                        ⚡ Smart Validation
                    </span>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )


# ============================================================================
# QUICK STATS
# ============================================================================

def render_quick_stats():
    """Render premium metric cards"""
    
    col1, col2, col3, col4 = st.columns(4, gap="small")
    
    with col1:
        st.markdown(
            f"""
            <div style="
                background-color: {COLORS['neutrals']['white']};
                border-radius: 16px;
                padding: 16px 8px;
                box-shadow: {SHADOWS['md']};
                text-align: center;
                border-top: 4px solid {COLORS['primary']['corporate_blue']};
                transition: transform 0.2s;
            ">
                <div style="font-size: 28px; font-weight: 600; color: {COLORS['primary']['deepest_navy']};">
                    {len(TEMPLATES)}
                </div>
                <div style="font-size: 14px; color: {COLORS['neutrals']['dark_gray']};">
                    📋 Available Templates
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    
    with col2:
        generated_count = st.session_state.get('generated_count', 0)
        st.markdown(
            f"""
            <div style="
                background-color: {COLORS['neutrals']['white']};
                border-radius: 16px;
                padding: 16px 8px;
                box-shadow: {SHADOWS['md']};
                text-align: center;
                border-top: 4px solid {COLORS['semantic']['success']};
                transition: transform 0.2s;
            ">
                <div style="font-size: 28px; font-weight: 600; color: {COLORS['primary']['deepest_navy']};">
                    {generated_count}
                </div>
                <div style="font-size: 14px; color: {COLORS['neutrals']['dark_gray']};">
                    ✅ Contracts Generated
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    
    with col3:
        active_warnings = st.session_state.get('active_warnings', 0)
        warning_color = COLORS['semantic']['warning'] if active_warnings > 0 else COLORS['neutrals']['medium_gray']
        st.markdown(
            f"""
            <div style="
                background-color: {COLORS['neutrals']['white']};
                border-radius: 16px;
                padding: 16px 8px;
                box-shadow: {SHADOWS['md']};
                text-align: center;
                border-top: 4px solid {warning_color};
                transition: transform 0.2s;
            ">
                <div style="font-size: 28px; font-weight: 600; color: {COLORS['primary']['deepest_navy']};">
                    {active_warnings}
                </div>
                <div style="font-size: 14px; color: {COLORS['neutrals']['dark_gray']};">
                    ⚠️ Active Warnings
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    
    with col4:
        last_generated = st.session_state.get('last_generated', 'Never')
        st.markdown(
            f"""
            <div style="
                background-color: {COLORS['neutrals']['white']};
                border-radius: 16px;
                padding: 16px 8px;
                box-shadow: {SHADOWS['md']};
                text-align: center;
                border-top: 4px solid {COLORS['semantic']['info']};
                transition: transform 0.2s;
            ">
                <div style="font-size: 20px; font-weight: 600; color: {COLORS['primary']['deepest_navy']};">
                    {last_generated}
                </div>
                <div style="font-size: 14px; color: {COLORS['neutrals']['dark_gray']};">
                    🕐 Last Generated
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )


# ============================================================================
# TEMPLATE SELECTOR
# ============================================================================

def render_template_selector():
    """Render premium template selector with descriptions and proper spacing"""
    
    st.markdown(
        f"""
        <div style="margin-bottom: 24px;">
            <h3 style="color: {COLORS['primary']['deepest_navy']}; font-size: 18px; font-weight: 600; margin: 0;">
                📋 Select Template
            </h3>
            <p style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; margin: 4px 0 0 0;">
                Choose from our curated collection of professional contract templates
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Template cards with 3-column grid
    template_options = TEMPLATES.copy()
    selected_template = st.session_state.get('selected_template', TEMPLATES[0])
    
    # Use 3 columns with consistent gap
    cols = st.columns(3, gap="medium")
    
    for idx, template_name in enumerate(template_options):
        col_idx = idx % 3
        with cols[col_idx]:
            is_selected = (selected_template == template_name)
            
            # Get template description
            description = TEMPLATE_DESCRIPTIONS.get(template_name, "")
            
            # Template icon mapping
            icons = {
                "NDA (Unilateral)": "🔒",
                "NDA (Mutual)": "🤝",
                "Master Services Agreement": "📋",
                "Independent Contractor Agreement": "👤",
                "Employment Offer Letter": "💼",
                "Software License Agreement": "💻"
            }
            icon = icons.get(template_name, "📄")
            
            # Template color mapping
            colors = {
                "NDA (Unilateral)": "#2C5F8A",
                "NDA (Mutual)": "#0D9488",
                "Master Services Agreement": "#D97706",
                "Independent Contractor Agreement": "#3B82F6",
                "Employment Offer Letter": "#0D9488",
                "Software License Agreement": "#7BA5C4"
            }
            border_color = colors.get(template_name, COLORS['primary']['corporate_blue'])
            
            # Determine styling
            if is_selected:
                bg_color = COLORS['primary']['ice_blue']
                text_color = COLORS['primary']['deepest_navy']
                desc_color = COLORS['neutrals']['dark_gray']
                shadow = SHADOWS['lg']
                border = border_color
                active_indicator = "✅"
                border_width = "3px"
            else:
                bg_color = COLORS['neutrals']['white']
                text_color = COLORS['neutrals']['dark_gray']
                desc_color = COLORS['neutrals']['medium_gray']
                shadow = SHADOWS['sm']
                border = COLORS['neutrals']['light_gray']
                active_indicator = ""
                border_width = "1.5px"
            
            # Render card with fixed height and proper padding
            st.markdown(
                f"""
                <div style="
                    background-color: {bg_color};
                    border: {border_width} solid {border};
                    border-radius: 16px;
                    padding: 20px 16px 16px 16px;
                    text-align: center;
                    box-shadow: {shadow};
                    height: 180px;
                    display: flex;
                    flex-direction: column;
                    justify-content: flex-start;
                    align-items: center;
                    position: relative;
                    transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1);
                    cursor: default;
                ">
                    <div style="font-size: 36px; margin-bottom: 6px; line-height: 1.2;">{icon}</div>
                    <div style="
                        font-size: 14px; 
                        font-weight: 600; 
                        color: {text_color};
                        line-height: 1.3;
                        margin-bottom: 4px;
                    ">{template_name}</div>
                    <div style="
                        font-size: 12px; 
                        color: {desc_color};
                        line-height: 1.4;
                        flex: 1;
                        display: flex;
                        align-items: center;
                    ">{description[:35]}{'...' if len(description) > 35 else ''}</div>
                    {f'<div style="position: absolute; top: 8px; right: 10px; font-size: 18px;">{active_indicator}</div>' if active_indicator else ''}
                </div>
                """,
                unsafe_allow_html=True
            )
            
            # Button with consistent spacing and styling
            btn_key = f"template_btn_{idx}"
            button_type = "primary" if is_selected else "secondary"
            
            # Add a small spacer between card and button
            st.markdown('<div style="height: 10px;"></div>', unsafe_allow_html=True)
            
            if st.button(
                f"Select {template_name}", 
                key=btn_key, 
                use_container_width=True, 
                type=button_type
            ):
                st.session_state.selected_template = template_name
                st.session_state.template_selected = True
                st.rerun()
            
            # Add extra spacing between items
            if idx < len(template_options) - 1:
                st.markdown('<div style="height: 4px;"></div>', unsafe_allow_html=True)


# ============================================================================
# PREMIUM FEATURES - COMPLETELY FIXED
# ============================================================================

def render_premium_features():
    """Render premium features and instructions - COMPLETELY FIXED HTML"""
    
    # Main container with gradient background
    st.markdown(
        f"""
        <div style="
            background: linear-gradient(135deg, {COLORS['primary']['ice_blue']} 0%, {COLORS['neutrals']['white']} 100%);
            border-radius: 16px;
            padding: 24px 20px;
            border: 1px solid {COLORS['neutrals']['light_gray']};
            height: 100%;
        ">
            <h4 style="color: {COLORS['primary']['deepest_navy']}; margin: 0 0 20px 0; font-size: 17px;">
                ✨ Premium Features
            </h4>
        """,
        unsafe_allow_html=True
    )
    
    # Feature 1: Smart Validation
    st.markdown(
        f"""
        <div style="margin-bottom: 16px;">
            <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 4px;">
                <span style="font-size: 20px;">🔒</span>
                <span style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; font-weight: 600;">Smart Validation</span>
                <span style="
                    background: {COLORS['semantic']['success_bg']};
                    color: {COLORS['semantic']['success']};
                    font-size: 10px;
                    padding: 2px 10px;
                    border-radius: 12px;
                    font-weight: 600;
                ">PRO</span>
            </div>
            <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px; margin: 0 0 0 34px; line-height: 1.4;">
                Real-time validation with actionable warnings and recommendations
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Feature 2: Smart Defaults
    st.markdown(
        f"""
        <div style="margin-bottom: 16px;">
            <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 4px;">
                <span style="font-size: 20px;">⚡</span>
                <span style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; font-weight: 600;">Smart Defaults</span>
                <span style="
                    background: {COLORS['semantic']['info_bg']};
                    color: {COLORS['semantic']['info']};
                    font-size: 10px;
                    padding: 2px 10px;
                    border-radius: 12px;
                    font-weight: 600;
                ">AI</span>
            </div>
            <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px; margin: 0 0 0 34px; line-height: 1.4;">
                Intelligent defaults based on contract type and industry standards
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Feature 3: Multiple Formats
    st.markdown(
        f"""
        <div style="margin-bottom: 16px;">
            <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 4px;">
                <span style="font-size: 20px;">📊</span>
                <span style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; font-weight: 600;">Multiple Formats</span>
            </div>
            <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px; margin: 0 0 0 34px; line-height: 1.4;">
                Export as TXT or professional DOCX with proper formatting
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Feature 4: Industry Standard Clauses
    st.markdown(
        f"""
        <div style="margin-bottom: 16px;">
            <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 4px;">
                <span style="font-size: 20px;">🎯</span>
                <span style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; font-weight: 600;">Standard Clauses</span>
            </div>
            <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px; margin: 0 0 0 34px; line-height: 1.4;">
                Expert-crafted clauses based on legal best practices
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Feature 5: Quick Generation
    st.markdown(
        f"""
        <div style="margin-bottom: 4px;">
            <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 4px;">
                <span style="font-size: 20px;">🚀</span>
                <span style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; font-weight: 600;">Quick Generation</span>
            </div>
            <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px; margin: 0 0 0 34px; line-height: 1.4;">
                Generate professional contracts in seconds, not hours
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Close the container div
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Quick tutorial - separate div with better styling
    st.markdown(
        f"""
        <div style="
            background: {COLORS['semantic']['info_bg']};
            border-radius: 12px;
            padding: 18px 16px;
            margin-top: 18px;
            border-left: 4px solid {COLORS['semantic']['info']};
        ">
            <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 10px;">
                <span style="font-size: 20px;">💡</span>
                <span style="color: {COLORS['primary']['deepest_navy']}; font-weight: 600; font-size: 15px;">How to Use</span>
            </div>
            <div style="color: {COLORS['neutrals']['dark_gray']}; font-size: 13px; padding-left: 20px; line-height: 1.8;">
                <div>1. Select a template from the cards above</div>
                <div>2. Fill in the required fields (marked with *)</div>
                <div>3. Review warnings and suggestions in real-time</div>
                <div>4. Click "Generate Contract" to preview</div>
                <div>5. Download your contract in TXT or DOCX format</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )


# ============================================================================
# SHARED UTILITY FUNCTIONS
# ============================================================================

def display_validation_warnings(warnings: List[Dict[str, Any]]):
    """Display validation warnings with premium styling"""
    
    if not warnings:
        return
    
    error_count = sum(1 for w in warnings if w.get('severity') == 'error')
    warning_count = sum(1 for w in warnings if w.get('severity') == 'warning')
    info_count = sum(1 for w in warnings if w.get('severity') == 'info')
    
    # Update session state for stats
    st.session_state.active_warnings = len(warnings)
    
    # Display summary badges
    badges_html = ""
    if error_count > 0:
        badges_html += f'<span style="background: #FEE2E2; color: #DC2626; padding: 4px 14px; border-radius: 12px; font-size: 12px; font-weight: 500; margin-right: 8px;">🔴 {error_count} Errors</span>'
    if warning_count > 0:
        badges_html += f'<span style="background: #FEF3C7; color: #D97706; padding: 4px 14px; border-radius: 12px; font-size: 12px; font-weight: 500; margin-right: 8px;">🟡 {warning_count} Warnings</span>'
    if info_count > 0:
        badges_html += f'<span style="background: #EFF6FF; color: #3B82F6; padding: 4px 14px; border-radius: 12px; font-size: 12px; font-weight: 500;">🔵 {info_count} Info</span>'
    
    if badges_html:
        st.markdown(f'<div style="margin: 12px 0;">{badges_html}</div>', unsafe_allow_html=True)
    
    # Display individual warnings
    for warning in warnings:
        severity = warning.get('severity', 'info')
        field = warning.get('field', '')
        message = warning.get('message', '')
        
        if severity == 'error':
            bg = '#FEE2E2'
            border = '#DC2626'
            icon = '❌'
            color = '#DC2626'
        elif severity == 'warning':
            bg = '#FEF3C7'
            border = '#D97706'
            icon = '⚠️'
            color = '#D97706'
        else:
            bg = '#EFF6FF'
            border = '#3B82F6'
            icon = 'ℹ️'
            color = '#3B82F6'
        
        st.markdown(
            f"""
            <div style="
                background: {bg};
                border-left: 4px solid {border};
                padding: 10px 14px;
                border-radius: 8px;
                margin-bottom: 6px;
                display: flex;
                align-items: start;
                gap: 8px;
            ">
                <span style="font-size: 16px;">{icon}</span>
                <div>
                    <span style="color: {color}; font-weight: 500; font-size: 13px;">
                        {field.replace('_', ' ').title()}
                    </span>
                    <span style="color: {COLORS['neutrals']['dark_gray']}; font-size: 13px; margin-left: 8px;">
                        {message}
                    </span>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )


def display_generated_contract():
    """Display the generated contract with premium styling"""
    
    if 'generated_contract' not in st.session_state:
        return
    
    contract_text = st.session_state.generated_contract
    template_name = st.session_state.get('last_template_used', 'Contract')
    
    st.divider()
    
    st.markdown(
        f"""
        <div style="margin: 16px 0;">
            <div style="display: flex; justify-content: space-between; align-items: center;">
                <h3 style="color: {COLORS['primary']['deepest_navy']}; font-size: 18px; font-weight: 600; margin: 0;">
                    📄 Generated {template_name}
                </h3>
                <span style="
                    background: {COLORS['semantic']['success_bg']};
                    color: {COLORS['semantic']['success']};
                    padding: 4px 14px;
                    border-radius: 12px;
                    font-size: 12px;
                    font-weight: 500;
                ">
                    ✅ Ready to Download
                </span>
            </div>
            <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px; margin: 4px 0 12px 0;">
                Review your contract below. You can copy it or download in your preferred format.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Contract preview with styling
    st.markdown(
        f"""
        <div style="
            background: {COLORS['neutrals']['white']};
            border: 1px solid {COLORS['neutrals']['light_gray']};
            border-radius: 12px;
            padding: 16px;
            max-height: 500px;
            overflow-y: auto;
            font-family: 'Courier New', monospace;
            font-size: 13px;
            line-height: 1.6;
            white-space: pre-wrap;
            color: {COLORS['neutrals']['dark_gray']};
        ">
            {contract_text}
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Download buttons
    st.markdown("### 📥 Download Options")
    
    col1, col2, col3 = st.columns(3, gap="small")
    
    with col1:
        if st.button("📋 Copy to Clipboard", key="copy_contract_btn", use_container_width=True):
            st.code(contract_text, language="text")
            st.success("✅ Contract copied! You can now copy it from the code block above.")
    
    with col2:
        st.download_button(
            label="📄 Download as TXT",
            data=contract_text,
            file_name=f"{template_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
            mime="text/plain",
            key="download_txt_btn",
            use_container_width=True
        )
    
    with col3:
        # Create real DOCX
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document()
            doc.add_heading(template_name, 0)
            
            # Add contract text with proper formatting
            for line in contract_text.split('\n'):
                if line.strip():
                    if line.strip().endswith(':') or line.strip().isupper():
                        doc.add_paragraph(line, style='Heading 2')
                    else:
                        doc.add_paragraph(line)
            
            # Save to bytes
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            st.download_button(
                label="📝 Download as DOCX",
                data=doc_bytes,
                file_name=f"{template_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx_btn",
                use_container_width=True
            )
        except ImportError:
            st.warning("⚠️ python-docx not installed. Install with: `pip install python-docx`")
            st.download_button(
                label="📝 Download as DOCX (Text Only)",
                data=contract_text,
                file_name=f"{template_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="text/plain",
                key="download_docx_fallback_btn",
                use_container_width=True
            )
    
    # Clear button - centered
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🗑️ Clear Generated Contract", key="clear_contract_btn", use_container_width=True):
            del st.session_state.generated_contract
            st.rerun()


# ============================================================================
# FOOTER TIPS
# ============================================================================

def render_footer_tips():
    """Render footer with pro tips"""
    
    st.divider()
    
    tips = [
        "💡 **Pro Tip:** All fields marked with * are required for contract generation",
        "🎯 **Smart Defaults:** We pre-fill common values to save you time",
        "⚠️ **Validation:** Pay attention to warnings - they help you avoid common pitfalls",
        "📊 **Multiple Formats:** Download as TXT for editing or DOCX for professional use",
        "🔒 **Privacy First:** All processing happens locally - your data never leaves your device"
    ]
    
    st.markdown("### 💡 Pro Tips")
    
    cols = st.columns(len(tips), gap="small")
    for col, tip in zip(cols, tips):
        with col:
            st.markdown(
                f"""
                <div style="
                    background: {COLORS['neutrals']['off_white']};
                    padding: 12px 8px;
                    border-radius: 8px;
                    text-align: center;
                    font-size: 12px;
                    color: {COLORS['neutrals']['dark_gray']};
                    height: 100%;
                    min-height: 60px;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    line-height: 1.4;
                ">
                    {tip}
                </div>
                """,
                unsafe_allow_html=True
            )


# ============================================================================
# INDIVIDUAL TEMPLATE RENDERERS
# ============================================================================

def render_nda_unilateral():
    """Render unilateral NDA form with premium UI"""
    
    st.markdown(
        f"""
        <div style="
            background: linear-gradient(135deg, {COLORS['primary']['ice_blue']} 0%, {COLORS['neutrals']['white']} 100%);
            border-radius: 16px;
            padding: 24px;
            border: 1px solid {COLORS['neutrals']['light_gray']};
            margin: 16px 0;
        ">
            <h3 style="color: {COLORS['primary']['deepest_navy']}; font-size: 18px; font-weight: 600; margin: 0;">
                🔒 NDA (Unilateral)
            </h3>
            <p style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; margin: 4px 0 0 0;">
                One-way non-disclosure agreement for protecting your confidential information
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Get template fields
    template = NDAUnilateralTemplate()
    fields = template.fields
    
    with st.form("nda_unilateral_form"):
        # Create two columns for fields
        col1, col2 = st.columns(2, gap="medium")
        
        field_values = {}
        
        with col1:
            # Left column fields
            for field_name in ['disclosing_party', 'receiving_party', 'purpose']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        with col2:
            # Right column fields
            for field_name in ['confidentiality_period', 'term_years', 'governing_law']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        st.divider()
        
        # Special terms
        st.markdown(
            f"""
            <div style="margin: 16px 0;">
                <h4 style="color: {COLORS['primary']['deepest_navy']}; font-size: 15px; font-weight: 600;">
                    ⚙️ Special Terms
                </h4>
                <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px;">
                    Customize your NDA with additional provisions
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        col1, col2 = st.columns(2, gap="medium")
        
        with col1:
            field_values['include_exclusions'] = st.checkbox(
                "Include Exclusions",
                value=True,
                help="Standard exclusions for publicly available information"
            )
        
        with col2:
            field_values['include_return'] = st.checkbox(
                "Include Return of Information",
                value=True,
                help="Require return or destruction of confidential information upon termination"
            )
        
        st.divider()
        
        # Real-time validation
        warnings = template.validate(field_values)
        display_validation_warnings(warnings)
        
        # Submit button - centered with better alignment
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            submitted = st.form_submit_button("📝 Generate Contract", use_container_width=True, type="primary")
        
        if submitted:
            # Check for critical errors
            critical_errors = [w for w in warnings if w.get('severity') == 'error']
            if critical_errors:
                st.error("❌ Please fix all errors before generating the contract")
            else:
                try:
                    contract = template.generate(field_values)
                    st.session_state.generated_contract = contract
                    st.session_state.last_template_used = "NDA (Unilateral)"
                    st.session_state.generated_count = st.session_state.get('generated_count', 0) + 1
                    st.session_state.last_generated = datetime.now().strftime('%H:%M')
                    st.success("✅ Contract generated successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error generating contract: {str(e)}")
    
    # Display generated contract if exists
    if 'generated_contract' in st.session_state and st.session_state.get('last_template_used') == "NDA (Unilateral)":
        display_generated_contract()


def render_nda_mutual():
    """Render mutual NDA form with premium UI"""
    
    st.markdown(
        f"""
        <div style="
            background: linear-gradient(135deg, {COLORS['primary']['ice_blue']} 0%, {COLORS['neutrals']['white']} 100%);
            border-radius: 16px;
            padding: 24px;
            border: 1px solid {COLORS['neutrals']['light_gray']};
            margin: 16px 0;
        ">
            <h3 style="color: {COLORS['primary']['deepest_navy']}; font-size: 18px; font-weight: 600; margin: 0;">
                🤝 NDA (Mutual)
            </h3>
            <p style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; margin: 4px 0 0 0;">
                Two-way non-disclosure agreement for mutual protection of confidential information
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Get template fields
    template = NDAMutualTemplate()
    fields = template.fields
    
    with st.form("nda_mutual_form"):
        # Create two columns for fields
        col1, col2 = st.columns(2, gap="medium")
        
        field_values = {}
        
        with col1:
            # Left column fields
            for field_name in ['party_a', 'party_b', 'purpose']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        with col2:
            # Right column fields
            for field_name in ['confidentiality_period', 'term_years', 'governing_law']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        st.divider()
        
        # Real-time validation
        warnings = template.validate(field_values)
        display_validation_warnings(warnings)
        
        # Submit button - centered with better alignment
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            submitted = st.form_submit_button("📝 Generate Contract", use_container_width=True, type="primary")
        
        if submitted:
            critical_errors = [w for w in warnings if w.get('severity') == 'error']
            if critical_errors:
                st.error("❌ Please fix all errors before generating the contract")
            else:
                try:
                    contract = template.generate(field_values)
                    st.session_state.generated_contract = contract
                    st.session_state.last_template_used = "NDA (Mutual)"
                    st.session_state.generated_count = st.session_state.get('generated_count', 0) + 1
                    st.session_state.last_generated = datetime.now().strftime('%H:%M')
                    st.success("✅ Contract generated successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error generating contract: {str(e)}")
    
    # Display generated contract if exists
    if 'generated_contract' in st.session_state and st.session_state.get('last_template_used') == "NDA (Mutual)":
        display_generated_contract()


def render_msa():
    """Render Master Services Agreement form with premium UI"""
    
    st.markdown(
        f"""
        <div style="
            background: linear-gradient(135deg, {COLORS['primary']['ice_blue']} 0%, {COLORS['neutrals']['white']} 100%);
            border-radius: 16px;
            padding: 24px;
            border: 1px solid {COLORS['neutrals']['light_gray']};
            margin: 16px 0;
        ">
            <h3 style="color: {COLORS['primary']['deepest_navy']}; font-size: 18px; font-weight: 600; margin: 0;">
                📋 Master Services Agreement
            </h3>
            <p style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; margin: 4px 0 0 0;">
                Comprehensive agreement for ongoing vendor and service provider relationships
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Get template fields
    template = MSATemplate()
    fields = template.fields
    
    with st.form("msa_form"):
        # Create two columns for fields
        col1, col2 = st.columns(2, gap="medium")
        
        field_values = {}
        
        with col1:
            # Left column fields
            for field_name in ['service_provider', 'client', 'services_description']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        with col2:
            # Right column fields
            for field_name in ['payment_terms', 'liability_cap', 'governing_law']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        st.divider()
        
        # Special terms
        st.markdown(
            f"""
            <div style="margin: 16px 0;">
                <h4 style="color: {COLORS['primary']['deepest_navy']}; font-size: 15px; font-weight: 600;">
                    ⚙️ Special Terms
                </h4>
                <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px;">
                    Customize your MSA with additional provisions
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        field_values['auto_renewal'] = st.checkbox(
            "Auto-Renewal",
            value=False,
            help="⚠️ Auto-renewal may lock you into the agreement"
        )
        
        st.divider()
        
        # Real-time validation
        warnings = template.validate(field_values)
        display_validation_warnings(warnings)
        
        # Submit button - centered with better alignment
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            submitted = st.form_submit_button("📝 Generate Contract", use_container_width=True, type="primary")
        
        if submitted:
            critical_errors = [w for w in warnings if w.get('severity') == 'error']
            if critical_errors:
                st.error("❌ Please fix all errors before generating the contract")
            else:
                try:
                    contract = template.generate(field_values)
                    st.session_state.generated_contract = contract
                    st.session_state.last_template_used = "Master Services Agreement"
                    st.session_state.generated_count = st.session_state.get('generated_count', 0) + 1
                    st.session_state.last_generated = datetime.now().strftime('%H:%M')
                    st.success("✅ Contract generated successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error generating contract: {str(e)}")
    
    # Display generated contract if exists
    if 'generated_contract' in st.session_state and st.session_state.get('last_template_used') == "Master Services Agreement":
        display_generated_contract()


def render_ica():
    """Render Independent Contractor Agreement form"""
    
    st.markdown(
        f"""
        <div style="
            background: linear-gradient(135deg, {COLORS['primary']['ice_blue']} 0%, {COLORS['neutrals']['white']} 100%);
            border-radius: 16px;
            padding: 24px;
            border: 1px solid {COLORS['neutrals']['light_gray']};
            margin: 16px 0;
        ">
            <h3 style="color: {COLORS['primary']['deepest_navy']}; font-size: 18px; font-weight: 600; margin: 0;">
                👤 Independent Contractor Agreement
            </h3>
            <p style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; margin: 4px 0 0 0;">
                Professional agreement for freelancers and independent contractors
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Get template fields
    template = IndependentContractorTemplate()
    fields = template.fields
    
    with st.form("ica_form"):
        # Create two columns for fields
        col1, col2 = st.columns(2, gap="medium")
        
        field_values = {}
        
        with col1:
            # Left column fields
            for field_name in ['company', 'contractor', 'services']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        with col2:
            # Right column fields
            for field_name in ['payment_rate', 'payment_terms', 'project_duration_months', 'governing_law']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        st.divider()
        
        # Real-time validation
        warnings = template.validate(field_values)
        display_validation_warnings(warnings)
        
        # Submit button - centered with better alignment
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            submitted = st.form_submit_button("📝 Generate Contract", use_container_width=True, type="primary")
        
        if submitted:
            critical_errors = [w for w in warnings if w.get('severity') == 'error']
            if critical_errors:
                st.error("❌ Please fix all errors before generating the contract")
            else:
                try:
                    contract = template.generate(field_values)
                    st.session_state.generated_contract = contract
                    st.session_state.last_template_used = "Independent Contractor Agreement"
                    st.session_state.generated_count = st.session_state.get('generated_count', 0) + 1
                    st.session_state.last_generated = datetime.now().strftime('%H:%M')
                    st.success("✅ Contract generated successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error generating contract: {str(e)}")
    
    # Display generated contract if exists
    if 'generated_contract' in st.session_state and st.session_state.get('last_template_used') == "Independent Contractor Agreement":
        display_generated_contract()


def render_employment():
    """Render Employment Offer Letter form"""
    
    st.markdown(
        f"""
        <div style="
            background: linear-gradient(135deg, {COLORS['primary']['ice_blue']} 0%, {COLORS['neutrals']['white']} 100%);
            border-radius: 16px;
            padding: 24px;
            border: 1px solid {COLORS['neutrals']['light_gray']};
            margin: 16px 0;
        ">
            <h3 style="color: {COLORS['primary']['deepest_navy']}; font-size: 18px; font-weight: 600; margin: 0;">
                💼 Employment Offer Letter
            </h3>
            <p style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; margin: 4px 0 0 0;">
                Professional offer letter for new employee hires
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Get template fields
    template = EmploymentOfferTemplate()
    fields = template.fields
    
    with st.form("employment_form"):
        # Create two columns for fields
        col1, col2 = st.columns(2, gap="medium")
        
        field_values = {}
        
        with col1:
            # Left column fields
            for field_name in ['company_name', 'candidate_name', 'position', 'start_date']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        with col2:
            # Right column fields
            for field_name in ['salary', 'employment_type', 'governing_law']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        st.divider()
        
        # Real-time validation
        warnings = template.validate(field_values)
        display_validation_warnings(warnings)
        
        # Submit button - centered with better alignment
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            submitted = st.form_submit_button("📝 Generate Contract", use_container_width=True, type="primary")
        
        if submitted:
            critical_errors = [w for w in warnings if w.get('severity') == 'error']
            if critical_errors:
                st.error("❌ Please fix all errors before generating the contract")
            else:
                try:
                    contract = template.generate(field_values)
                    st.session_state.generated_contract = contract
                    st.session_state.last_template_used = "Employment Offer Letter"
                    st.session_state.generated_count = st.session_state.get('generated_count', 0) + 1
                    st.session_state.last_generated = datetime.now().strftime('%H:%M')
                    st.success("✅ Contract generated successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error generating contract: {str(e)}")
    
    # Display generated contract if exists
    if 'generated_contract' in st.session_state and st.session_state.get('last_template_used') == "Employment Offer Letter":
        display_generated_contract()


def render_software_license():
    """Render Software License Agreement form"""
    
    st.markdown(
        f"""
        <div style="
            background: linear-gradient(135deg, {COLORS['primary']['ice_blue']} 0%, {COLORS['neutrals']['white']} 100%);
            border-radius: 16px;
            padding: 24px;
            border: 1px solid {COLORS['neutrals']['light_gray']};
            margin: 16px 0;
        ">
            <h3 style="color: {COLORS['primary']['deepest_navy']}; font-size: 18px; font-weight: 600; margin: 0;">
                💻 Software License Agreement
            </h3>
            <p style="color: {COLORS['neutrals']['dark_gray']}; font-size: 14px; margin: 4px 0 0 0;">
                Professional software licensing and distribution agreement
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Get template fields
    template = SoftwareLicenseTemplate()
    fields = template.fields
    
    with st.form("software_license_form"):
        # Create two columns for fields
        col1, col2 = st.columns(2, gap="medium")
        
        field_values = {}
        
        with col1:
            # Left column fields
            for field_name in ['licensor', 'licensee', 'software_name', 'license_type']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        with col2:
            # Right column fields
            for field_name in ['license_fee', 'payment_terms', 'governing_law']:
                field_config = fields[field_name]
                field_values[field_name] = render_form_field(field_name, field_config)
        
        st.divider()
        
        # Special terms
        st.markdown(
            f"""
            <div style="margin: 16px 0;">
                <h4 style="color: {COLORS['primary']['deepest_navy']}; font-size: 15px; font-weight: 600;">
                    ⚙️ License Terms
                </h4>
                <p style="color: {COLORS['neutrals']['medium_gray']}; font-size: 13px;">
                    Customize your software license terms
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        field_values['maintenance_support'] = st.checkbox(
            "Include Maintenance & Support",
            value=True,
            help="Add maintenance and support provisions"
        )
        
        st.divider()
        
        # Real-time validation
        warnings = template.validate(field_values)
        display_validation_warnings(warnings)
        
        # Submit button - centered with better alignment
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            submitted = st.form_submit_button("📝 Generate Contract", use_container_width=True, type="primary")
        
        if submitted:
            critical_errors = [w for w in warnings if w.get('severity') == 'error']
            if critical_errors:
                st.error("❌ Please fix all errors before generating the contract")
            else:
                try:
                    contract = template.generate(field_values)
                    st.session_state.generated_contract = contract
                    st.session_state.last_template_used = "Software License Agreement"
                    st.session_state.generated_count = st.session_state.get('generated_count', 0) + 1
                    st.session_state.last_generated = datetime.now().strftime('%H:%M')
                    st.success("✅ Contract generated successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Error generating contract: {str(e)}")
    
    # Display generated contract if exists
    if 'generated_contract' in st.session_state and st.session_state.get('last_template_used') == "Software License Agreement":
        display_generated_contract()


# ============================================================================
# FORM FIELD RENDERER
# ============================================================================

def render_form_field(field_name: str, field_config: Dict[str, Any]):
    """Render a form field based on its configuration"""
    
    field_type = field_config.get('type', 'text')
    label = field_config.get('label', field_name.replace('_', ' ').title())
    required = field_config.get('required', False)
    placeholder = field_config.get('placeholder', '')
    default = field_config.get('default', '')
    min_val = field_config.get('min')
    max_val = field_config.get('max')
    options = field_config.get('options', [])
    help_text = field_config.get('help', '')
    
    # Add required indicator
    display_label = f"{label} {'*' if required else ''}"
    
    if field_type == 'text':
        return st.text_input(
            display_label,
            value=default if default else '',
            placeholder=placeholder,
            help=help_text,
            key=f"field_{field_name}"
        )
    
    elif field_type == 'textarea':
        return st.text_area(
            display_label,
            value=default if default else '',
            placeholder=placeholder,
            help=help_text,
            key=f"field_{field_name}",
            height=100
        )
    
    elif field_type == 'number':
        return st.number_input(
            display_label,
            min_value=min_val,
            max_value=max_val,
            value=int(default) if default else 0,
            help=help_text,
            key=f"field_{field_name}"
        )
    
    elif field_type == 'select':
        return st.selectbox(
            display_label,
            options=options,
            index=options.index(default) if default in options else 0,
            help=help_text,
            key=f"field_{field_name}"
        )
    
    elif field_type == 'checkbox':
        return st.checkbox(
            display_label,
            value=default if isinstance(default, bool) else False,
            help=help_text,
            key=f"field_{field_name}"
        )
    
    else:
        return st.text_input(
            display_label,
            value=default if default else '',
            placeholder=placeholder,
            help=help_text,
            key=f"field_{field_name}"
        )