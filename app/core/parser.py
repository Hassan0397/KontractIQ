"""
KontractIQ - Document Parser
Handles PDF, DOCX, and TXT file parsing
"""

import io
from typing import Tuple, Optional, List, Dict, Any
import PyPDF2
import pdfplumber
import docx
from ..utils.constants import LIMITS
from ..utils.helpers import is_scanned_pdf


class DocumentParser:
    """Parse documents and extract text - OPTIMIZED"""
    
    @staticmethod
    def parse_pdf(file_content: bytes) -> Tuple[str, int, bool]:
        """Parse PDF file and extract text - FAST"""
        text = ""
        pages = 0
        is_scanned = False
        
        try:
            # Try PyPDF2 first (faster)
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            pages = len(pdf_reader.pages)
            
            # Limit pages to prevent memory issues
            max_pages = min(pages, LIMITS["max_pages"])
            
            for i in range(max_pages):
                try:
                    page_text = pdf_reader.pages[i].extract_text()
                    if page_text:
                        text += page_text + "\n"
                except:
                    continue
            
            # Check if scanned
            if is_scanned_pdf(text, LIMITS["min_text_length"]):
                is_scanned = True
                # Try pdfplumber as fallback (slower but better)
                try:
                    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                        pdf_text = ""
                        max_pages = min(len(pdf.pages), LIMITS["max_pages"])
                        for i in range(max_pages):
                            try:
                                page_text = pdf.pages[i].extract_text()
                                if page_text:
                                    pdf_text += page_text + "\n"
                            except:
                                continue
                        if not is_scanned_pdf(pdf_text, LIMITS["min_text_length"]):
                            text = pdf_text
                            is_scanned = False
                except:
                    pass
                
        except Exception as e:
            raise Exception(f"Failed to parse PDF: {str(e)}")
        
        # Limit text length
        if len(text) > 100000:
            text = text[:100000]
        
        return text, pages, is_scanned
    
    @staticmethod
    def parse_docx(file_content: bytes) -> Tuple[str, int, bool]:
        """Parse DOCX file and extract text - FAST"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            pages = len(paragraphs) // 30 + 1
            
            # Limit text
            if len(text) > 100000:
                text = text[:100000]
            
            is_scanned = is_scanned_pdf(text, LIMITS["min_text_length"])
            return text, pages, is_scanned
        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_txt(file_content: bytes) -> Tuple[str, int, bool]:
        """Parse TXT file and extract text - FAST"""
        try:
            text = file_content.decode('utf-8')
            pages = len(text) // 2000 + 1
            
            # Limit text
            if len(text) > 100000:
                text = text[:100000]
            
            is_scanned = False
            return text, pages, is_scanned
        except UnicodeDecodeError:
            try:
                text = file_content.decode('latin-1')
                pages = len(text) // 2000 + 1
                if len(text) > 100000:
                    text = text[:100000]
                is_scanned = False
                return text, pages, is_scanned
            except Exception as e:
                raise Exception(f"Failed to parse TXT: {str(e)}")
    
    @staticmethod
    def parse_file(file_content: bytes, file_type: str) -> Tuple[str, int, bool]:
        """Parse file based on type"""
        if file_type == "pdf":
            return DocumentParser.parse_pdf(file_content)
        elif file_type == "docx":
            return DocumentParser.parse_docx(file_content)
        elif file_type == "txt":
            return DocumentParser.parse_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def extract_metadata(file_content: bytes, file_type: str) -> Dict[str, Any]:
        """
        Extract metadata from file - OPTIMIZED
        
        Args:
            file_content: File content as bytes
            file_type: File type (pdf, docx, txt)
            
        Returns:
            Dict: Metadata including title, author, creation_date, etc.
        """
        metadata = {
            'title': None,
            'author': None,
            'creation_date': None,
            'modification_date': None,
            'page_count': 0,
            'word_count': 0,
            'character_count': 0
        }
        
        try:
            if file_type == 'pdf':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                metadata['page_count'] = len(pdf_reader.pages)
                
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', None)
                    metadata['author'] = pdf_reader.metadata.get('/Author', None)
                    metadata['creation_date'] = pdf_reader.metadata.get('/CreationDate', None)
                    metadata['modification_date'] = pdf_reader.metadata.get('/ModDate', None)
            
            elif file_type == 'docx':
                doc = docx.Document(io.BytesIO(file_content))
                if doc.core_properties:
                    metadata['title'] = doc.core_properties.title
                    metadata['author'] = doc.core_properties.author
                    metadata['creation_date'] = doc.core_properties.created
                    metadata['modification_date'] = doc.core_properties.modified
                    metadata['page_count'] = len(doc.paragraphs) // 30 + 1
            
            elif file_type == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
                metadata['word_count'] = len(text.split())
                metadata['character_count'] = len(text)
                metadata['page_count'] = len(text) // 2000 + 1
                
        except Exception:
            pass
        
        return metadata
    
    @staticmethod
    def extract_batch_metadata(files: List[Tuple[bytes, str]]) -> List[Dict[str, Any]]:
        """
        Extract metadata from multiple files
        
        Args:
            files: List of (file_content, file_type) tuples
            
        Returns:
            List of metadata dictionaries
        """
        results = []
        for file_content, file_type in files:
            metadata = DocumentParser.extract_metadata(file_content, file_type)
            results.append(metadata)
        return results